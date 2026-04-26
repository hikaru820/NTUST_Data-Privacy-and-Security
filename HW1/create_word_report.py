from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_report():
    doc = Document()

    # 設定中文字體：微軟正黑體
    doc.styles['Normal'].font.name = u'Microsoft JhengHei'  # type: ignore[union-attr]
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]
    doc.styles['Normal'].font.size = Pt(12)  # type: ignore[union-attr]

    # 標題
    title = doc.add_heading('期中作業 midterm-HW1 - 差異隱私合成數據生成與效能對比分析報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = u'Microsoft JhengHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]

    # 1. 實作任務與研究動機
    h1 = doc.add_heading('1. 實作任務與研究動機 (Task Overview)', level=1)
    for run in h1.runs: run.font.name = u'Microsoft JhengHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]
    
    doc.add_paragraph(
        "在現代數據分享情境中，如何在保證「個體隱私不被洩漏」的前提下，最大化保留資料的「機器學習分析價值 (Utility)」，"
        "是一項極具挑戰的任務。本報告針對 HW1 實作之 K-匿名性 (K-Anonymity) 進行延伸比對，"
        "引進嚴謹的數學定義——差異隱私 (Differential Privacy, DP)。我們透過生成差異隱私合成數據 (DP Synthetic Data) "
        "作為機器學習的新世代訓練集，並探討不同的隱私預算 (Privacy Budget, ε) 投入時，"
        "對多層感知器 (MLP) 模型預測效能的實際衝擊。"
    )

    # 2. 工具與演算法選擇
    h2 = doc.add_heading('2. 工具與演算法選擇 (Tool & Algorithm Selection)', level=1)
    for run in h2.runs: run.font.name = u'Microsoft JhengHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]

    doc.add_heading('2.1 採用開源工具', level=2)
    doc.add_paragraph("本研究不採直覺的「隨機加噪」，而是導入了國際知名套件 DataSynthesizer。")

    doc.add_heading('2.2 核心演算法：PrivBayes (NIST 競賽獲獎演算法)', level=2)
    p = doc.add_paragraph()
    p.add_run("作法與理由：").bold = True
    p.add_run("DataSynthesizer 內建之 Correlated Attribute Mode 實作了著名的 PrivBayes 演算法。對於如 Adult Income 這種具高度複雜連帶關係的表格型資料 (Tabular Data)，如果只對單一欄位獨立加入雜訊，會完全摧毀機器學習最重視的「特徵關聯」（例如：教育程度與收入的強耦合）。")
    
    p = doc.add_paragraph()
    p.add_run("優勢：").bold = True
    p.add_run("PrivBayes 會先透過學習真實資料，建構出多維度的貝氏網路 (Bayesian Network)。演算法在計算條件機率與網路權重時，會精確投入拉普拉斯噪音 (Laplace Noise) 來滿足 ε-DP 保護。如此產生的假資料，不僅從根本上切斷了與「真實個體」的直接連結，又能在總體統計分佈上，高度保留原始資料的關聯特性。")

    # 3. 實作流程與實踐步驟
    h3 = doc.add_heading('3. 實作流程與實踐步驟 (Methodology)', level=1)
    for run in h3.runs: run.font.name = u'Microsoft JhengHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]

    doc.add_paragraph("為確保對比實驗的嚴謹性 (Controlled Experiment)，本次實驗架構嚴格遵循「控制變因」原則：")
    doc.add_paragraph("1. 基準對齊 (Baseline Alignment)：延續 HW1 階段已清洗之 adult_cleaned.csv 作為唯一基準資料 (共 45,222 筆)。確保特徵刪除等邏輯與 HW1 完全一致。")
    doc.add_paragraph("2. 隱私預算分配 (Epsilon Setup)：設計五個級別的差異隱私保護強度，"
                       "分別為 ε = 0.1（極高度保護）、ε = 0.5、ε = 1.0（中度保護）、ε = 5.0、ε = 10.0（輕度保護/高實用性）。"
                       "利用 PrivBayes 動態生成五組等量（各 45,222 筆）的合成數據集，"
                       "涵蓋從強隱私到高效用的完整保護力度範圍。")
    doc.add_paragraph("3. 模型訓練與復現 (Model Training)：將這三組「完全無真實隱私風險」的合成數據，重新拋入 HW1 架構之 MLP 分類器 (兩層隱藏層 64-32，Adam 優化器) 中擬合，並由未被污染的 Test-set 驗證其泛化能力。")

    # 4. 比較結果分析
    h4 = doc.add_heading('4. 比較結果分析 (Results Analysis)', level=1)
    for run in h4.runs: run.font.name = u'Microsoft JhengHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]
    
    doc.add_paragraph(
        "以下表格彙整所有實驗結果，包含 SVM 與 MLP 兩種分類器在各隱私預算 ε 下的表現，"
        "以及 TVD（Total Variation Distance，統計相似性指標，越低代表合成數據越接近真實分佈）。"
        "所有 DP 實驗皆以 TSTR（Train on Synthetic, Test on Real）方式評估。"
    )

    # 加入表格（SVM + MLP 完整比較，含 Runtime 與 TVD）
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方法'
    hdr_cells[1].text = '參數'
    hdr_cells[2].text = 'Runtime (s)'
    hdr_cells[3].text = 'TVD ↓'
    hdr_cells[4].text = 'SVM AUC'
    hdr_cells[5].text = 'SVM Acc'
    hdr_cells[6].text = 'MLP AUC'
    hdr_cells[7].text = 'MLP Acc'

    records = [
        ('無保護 (基準)',     'Original',          '—',    '—',      '0.9200', '0.8580', '0.9166', '0.8628'),
        ('K-Anonymity (HW1)', 'K = 10',            '—',    '—',      '0.9085', '0.8517', '0.9120', '0.8579'),
        ('DP Synthetic',      'ε = 10.0 (低保護)', '40.6', '0.1884', '0.8683', '0.8253', '0.8881', '0.8441'),
        ('DP Synthetic',      'ε = 5.0',           '38.8', '0.1902', '0.8654', '0.8216', '0.8918', '0.8448'),
        ('DP Synthetic',      'ε = 1.0',           '39.6', '0.2393', '0.8448', '0.8153', '0.8672', '0.8263'),
        ('DP Synthetic',      'ε = 0.5',           '40.0', '0.2401', '0.7098', '0.7563', '0.8096', '0.7677'),
        ('DP Synthetic',      'ε = 0.1 (高保護)',  '43.3', '0.3338', '0.7948', '0.7347', '0.8162', '0.8014'),
    ]
    for m, p, rt, tvd, sauc, sacc, mauc, macc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = p
        row_cells[2].text = rt
        row_cells[3].text = tvd
        row_cells[4].text = sauc
        row_cells[5].text = sacc
        row_cells[6].text = mauc
        row_cells[7].text = macc

    # 5. 自動化對比圖表解讀指南
    h5 = doc.add_heading('5. 自動化對比圖表解讀指南 (How to Interpret the Visualizations)', level=1)
    for run in h5.runs: run.font.name = u'Microsoft JhengHei'; run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Microsoft JhengHei')  # type: ignore[union-attr]
    
    doc.add_paragraph("在這份雙軸柱狀圖中（圖表為 HW1_Midterm_Comparison.png），我們向團隊展示了「隱私-效能權衡 (Privacy-Utility Trade-off)」的強大證據。報告重點如下：")
    
    doc.add_paragraph("1. 軸線定義與視覺觀察：", style='List Number')
    doc.add_paragraph("X 軸代表不同的防護力度設計；Y 軸則顯示該防護力度下的模型表現（左圖為 AUC 模型鑑別度，右圖為 Accuracy 預測準確度）。")

    doc.add_paragraph("2. K-匿名性 (HW1) vs. DP合成表現：", style='List Number')
    doc.add_paragraph("從圖中左側群集可見，K-Anonymity 的衰退曲線非常平緩。因為它的本質是「模糊化 (Generalization)」，數值仍反映真實世界的分佈邊界，故其 AUC 仍能死守在 90% 以上。相比之下，DP 合成數據 (淺藍色區塊) 則是採用了更嚴苛的數學防護，直接從源頭抹除真實資料個體。")

    doc.add_paragraph("3. 差異隱私內部的「長尾衰退效應」解讀 (核心報告亮點)：", style='List Number')
    p = doc.add_paragraph()
    p.add_run("適度保護 (ε=10.0 / ε=5.0)：").bold = True
    p.add_run("當隱私預算放寬至 ε=10.0 時，MLP AUC 達 88.8%、SVM AUC 達 86.8%，"
              "Accuracy 分別為 84.4% 與 82.5%。TVD 僅 0.1884，"
              "代表合成數據的統計分佈高度貼近真實資料。"
              "此效能逼近 HW1 K-Anonymity (K=10) 的 MLP AUC 91.2%，"
              "證明 DP 合成數據足以取代高敏感原始資料供第三方研發使用。")

    p = doc.add_paragraph()
    p.add_run("極端防禦的代價 (ε=0.1)：").bold = True
    p.add_run("當 ε=0.1 時，TVD 上升至 0.3338，代表合成資料的分佈已與真實資料有明顯落差。"
              "SVM AUC 降至 79.5%、MLP AUC 降至 81.6%，Accuracy 分別為 73.5% 與 80.1%。"
              "為滿足嚴苛的差異隱私數學假設，PrivBayes 被迫加入大量 Laplace 噪音，"
              "破壞特徵間的關聯結構，印證了「絕對的隱私保護等同於部分資料價值的犧牲」。")

    # 插入圖片
    img_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ML part", "HW1_Midterm_Comparison.png")
    if os.path.exists(img_path):
        doc.add_paragraph("附圖：模型評估結果之柱狀圖對比。")
        doc.add_picture(img_path, width=Pt(400))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("\n[圖片 HW1_Midterm_Comparison.png 請在此處插入]")

    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Midterm_Differentially Private Synthetic Data-HW1.docx")
    doc.save(output_path)
    print("Docx created:", output_path)

if __name__ == '__main__':
    create_report()
